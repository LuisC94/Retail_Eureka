import os
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_diagram(filename, boxes, arrows, title):
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis('off')
    
    # Draw boxes
    for box in boxes:
        x, y, w, h, text, color = box
        rect = patches.Rectangle((x, y), w, h, linewidth=1, edgecolor='black', facecolor=color)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, text, horizontalalignment='center', verticalalignment='center', fontsize=9, weight='bold')
    
    # Draw arrows
    for arrow in arrows:
        x1, y1, x2, y2, text = arrow
        ax.annotate(text, xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(facecolor='black', shrink=0.05, width=1.5, headwidth=8),
                    fontsize=8, horizontalalignment='center', verticalalignment='center', color='blue')
                    
    plt.title(title, fontsize=14, weight='bold')
    plt.tight_layout()
    plt.savefig(filename, dpi=300)
    plt.close()

# Define diagrams
# Diagram 1: Broker (Redesigned for clarity)
b1 = [
    (2, 40, 15, 20, "Agente\nQ-Learning", "orange"),
    (25, 40, 20, 20, "Environment\n(Middleware)", "lightgreen"),
    (55, 30, 15, 40, "Message\nBroker\n(Eventos)", "lightblue"),
    (80, 70, 18, 15, "Plataforma/BD\n(Stock/Qualidade)", "lightgray"),
    (80, 42.5, 18, 15, "Order Agent\n(Encomendas)", "lightcoral"),
    (80, 15, 18, 15, "LIACC\n(Transporte)", "lightyellow"),
]
a1 = [
    (17, 52, 25, 52, "Action"),
    (25, 48, 17, 48, "State / Reward"),
    
    (45, 52, 55, 52, "Publica Ação"),
    (55, 48, 45, 48, "Ouve Eventos"),
    
    (70, 77, 80, 77, "Requerimento"),
    (80, 73, 70, 73, "Resposta/Evento"),
    
    (70, 50, 80, 50, "Envia/Lê"),
    (80, 46, 70, 46, "Pub/Sub"),
    
    (70, 22, 80, 22, "Rotas/Tempos"),
    (80, 18, 70, 18, "Eventos")
]
create_diagram("diag1_broker.png", b1, a1, "Opção 1: Arquitetura Message Broker + Environment")

# Diagram 2: P2P
b2 = [
    (40, 40, 25, 20, "Agente de Stock", "orange"),
    (5, 75, 20, 15, "Plataforma/BD", "lightgray"),
    (75, 75, 20, 15, "Order Agent", "lightcoral"),
    (40, 10, 20, 15, "LIACC", "lightyellow")
]
a2 = [
    (45, 60, 20, 75, "HTTP GET"),
    (55, 60, 80, 75, "API Call"),
    (50, 40, 50, 25, "API Call")
]
create_diagram("diag2_p2p.png", b2, a2, "Outras Opções: Comunicação Direta (P2P)")

# Diagram 3: Shared DB
b3 = [
    (35, 40, 30, 20, "Base de Dados\nCompartilhada", "lightblue"),
    (5, 75, 20, 15, "Agente de Stock", "orange"),
    (75, 75, 20, 15, "Order Agent", "lightcoral"),
    (40, 10, 20, 15, "LIACC", "lightyellow"),
]
a3 = [
    (20, 75, 40, 60, "Leitura/Escrita"),
    (80, 75, 60, 60, ""),
    (50, 25, 50, 40, "")
]
create_diagram("diag3_db.png", b3, a3, "Outras Opções: Base de Dados Partilhada")

doc = Document()
# Title formatting
title = doc.add_heading('Integração do Agente Q-Learning (Stock Management)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('1. Opção 1: Arquitetura Mediada (Message Broker + Environment) [Recomendada]', level=1)
doc.add_paragraph('Esta é a abordagem mais recomendada, pois garante que o agente atue de forma segura e rápida sem acoplar diretamente o cérebro (Agente IA) com o corpo (Plataforma, BD, etc). O sistema está dividido em três peças centrais:')
doc.add_picture('diag1_broker.png', width=Inches(6.5))

doc.add_heading('Explicação Detalhada da Opção 1', level=2)
doc.add_paragraph(' A. O "Cérebro" (Agente Q-Learning): Fica completamente cego e isolado. Ele não precisa de chaves de base de dados, nem sabe se os dados vieram do LIACC ou do Order Agent. A sua única função na vida é receber um vetor de Estado (números), fazer o seu cálculo matricial, e devolver uma Ação (também em formato matemático).', style='List Bullet')
doc.add_paragraph(' B. O "Assistente / Tradutor" (Environment / Middleware): É a peça que o Agente consulta. Quando o Agente quer agir, ele interage com o Environment. O Environment é o software programado na plataforma que sabe ir à Base de Dados buscar a capacidade ou o custo hordcoded e cruza essa informação com os eventos em tempo real lidos do Message Broker.', style='List Bullet')
doc.add_paragraph(' C. Os "Mensageiros" (Message Broker - ex: RabbitMQ / Kafka): Funciona como o correio central de eventos da plataforma. Em vez de o Order Agent ligar diretamente ao Stock Agent para avisar de uma encomenda, ele simplesmente "publica" numa caixa (Tópico) a mensagem de nova encomenda. O Environment, que está subscrito (a escutar) essa caixa, lê a mensagem e incorpora a informação de imediato.', style='List Bullet')

doc.add_paragraph('Resumo do Fluxo Passo a Passo:', style='Intense Quote')
doc.add_paragraph('1. Os vários sistemas (Plataforma, LIACC, e outros agentes) enviam atualizações independentes para o Message Broker de forma assíncrona.\n2. O Environment (Middleware) está à escuta. Assim que recolhe as métricas e junta com as variáveis hardcoded da Base de Dados, constrói a representação do mundo atual.\n3. O Agente Q-Learning pede esse estado ao Environment, analisa-o e devolve uma decisão (Action).\n4. O Environment traduz essa decisão, executa os passos na plataforma, afere o custo do resultado da mesma, envia a Recompensa (Reward) final para o Agente poder aprender, e publica no Broker o resultado da operação.')


doc.add_heading('2. Análise de Alternativas', level=1)
doc.add_paragraph('Existem outras abordagens teóricas em arquiteturas de microsserviços e multi-agente, aqui detalhadas sucintamente e evidenciando a fragilidade comparativa perante a Opção 1.')

doc.add_heading('2.1 Comunicação Direta Ponto-a-Ponto (P2P)', level=2)
doc.add_paragraph('O Agente tem conhecimento de todos os microsserviços e faz chamadas diretas REST ou conexões RPC para toda a parte.')
doc.add_picture('diag2_p2p.png', width=Inches(6.5))
doc.add_paragraph('Vantagem: Desenvolvimento inicial intuitivo; respostas imediatamente síncronas.', style='List Bullet')
doc.add_paragraph('Desvantagem: Criações de arquiteturas complexas tipo "Esparguete", que se tornam impossíveis de depurar. Redução drástica da robustez; se o LIACC tiver uma pequena falha de rede de 2 segundos, o Agente de Stock falha o seu ciclo inteiro (Timeouts em catadupa).', style='List Bullet')

doc.add_heading('2.2 Base de Dados Partilhada (Shared Database)', level=2)
doc.add_paragraph('Todos os agentes escrevem as suas previsões e lêem continuamente da mesma base de dados transacional.')
doc.add_picture('diag3_db.png', width=Inches(6.5))
doc.add_paragraph('Vantagem: Toda a gente acede ao mesmo local providenciando uma "Única Fonte de Verdade" simplificada de consultar via SQL.', style='List Bullet')
doc.add_paragraph('Desvantagem: Destrói a escalabilidade do sistema transacional. Os agentes têm de consultar de x em x segundos a DB num processo de Polling, o que resulta num pico exaustivo de recursos, transformando o "Core" da Plataforma no gargalo de toda a operação (Bottleneck grave).', style='List Bullet')

doc.add_heading('2.3 Padrão "Blackboard" ou API Gateway Unificado', level=2)
doc.add_paragraph('Um ambiente onde há um "Quadro na Parede" colaborativo ou um API Gateway onde tudo passa, e os agentes vão tirando e pondo conclusões à vez de forma controlada.')
doc.add_paragraph('Vantagem: Centraliza as permissões de acesso num ponto cego perfeitamente isolado.', style='List Bullet')
doc.add_paragraph('Desvantagem: Delega demasiada complexidade estrutural e de processamento ao Quadro ou ao Gateway da plataforma (A plataforma necessitará de desenhar a malha conectiva e arcar ativamente o processamento que deveria ser distribuído por cada agente).', style='List Bullet')


doc.add_heading('3. Conclusões e Justificação da Opção 1', level=1)
doc.add_paragraph('A Opção 1 ganha inequivocamente em arquiteturas que integram Inteligência Artificial, pelas seguintes justificações vitais:')

doc.add_paragraph(' Encaixe perfeito com Q-Learning: A abstração puramente matricial protege o cérebro lógico de problemas transacionais do mundo do software (timeout, HTTP erros), mantendo o ciclo Estado/Ação/Recompensa 100% puro para a matemática.', style='List Bullet')
doc.add_paragraph(' Desacoplamento e Escalabilidade (Plug-and-Play): Adicionar um novo agente amanhã, não requer atualizar o Agente de Stock, apenas mandá-lo subscrever a o canal respetivo. Sistemas assíncronos não bloqueiam uns atrás dos outros.', style='List Bullet')
doc.add_paragraph(' Segurança por Isolamento (Zero-Trust): Com o broker limitas os tópicos que os agentes podem ler e nunca há partilha de tokens da Base de Dados relacional, mantendo o Core da BD intocável da I.A.', style='List Bullet')

try:
    doc.save('Arquitetura_Agentes_Q_Learning.docx')
    print("Document regenerated successfully.")
except Exception as e:
    print(f"Error saving document: {e}")
    exit(1)

# Cleanup
try:
    os.remove("diag1_broker.png")
    os.remove("diag2_p2p.png")
    os.remove("diag3_db.png")
except:
    pass
