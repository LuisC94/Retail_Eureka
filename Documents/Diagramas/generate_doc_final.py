from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_document():
    document = Document()

    # Title
    title = document.add_heading('Esquema da Base de Dados e Funcionamento do Sistema', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph('Este documento descreve a estrutura da base de dados (em teoria) e o fluxo de dados do sistema para apoio à sua apresentação.')

    # Section 1
    document.add_heading('1. Visão Geral (Conceito)', level=1)
    p = document.add_paragraph('O sistema Retail-Eureka funciona como uma plataforma integrada que rastreia todo o ciclo de vida do produto agrícola, desde o planeamento da plantação até à entrega ao retalhista/consumidor.')
    
    p = document.add_paragraph('A arquitetura de dados divide-se em 3 grandes blocos:')
    p.style = 'List Number'
    document.add_paragraph('Produção (Upstream): Gestão de pomares, registo de operações (rega, fertilização) e colheitas.', style='List Bullet')
    document.add_paragraph('Mercado e Logística (Midstream): Publicação de ofertas/procuras, negociação e rastreio de transporte.', style='List Bullet')
    document.add_paragraph('Imutabilidade (Blockchain): Registo paralelo e seguro dos eventos críticos para garantir a rastreabilidade "Do Prado ao Prato".', style='List Bullet')

    # Section 2
    document.add_heading('2. Diagrama de Entidades (ERD Simplificado)', level=1)
    document.add_paragraph('O diagrama abaixo ilustra as principais relações entre as tabelas da base de dados.')
    
    # Insert Image
    if os.path.exists("schema_diagram.png"):
        document.add_picture('schema_diagram.png', width=Inches(6.0))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        document.add_paragraph("[ERRO: Imagem do diagrama não encontrada]")

    # Section 3
    document.add_heading('3. Explicação Funcional "Como Funciona"', level=1)

    document.add_heading('A. O Ciclo de Produção (O Produtor)', level=2)
    document.add_paragraph('1. Registo do Pomar (PlantationPlan):', style='List Number')
    document.add_paragraph('O produtor define o seu "ativo" principal. Regista a localização, a área, o tipo de solo e que culturas (PlantationCrop) estão plantadas.', style='List Continue')
    
    document.add_paragraph('2. Diário de Operações (PlantationEvent):', style='List Number')
    document.add_paragraph('Ao longo do ano, o produtor regista o que faz no pomar. Cada evento tem uma data e um tipo específico (Fertilização, Rega, Poda). A BD tem tabelas específicas para cada detalhe (ex: FertilizerSyntheticData).', style='List Continue')

    document.add_paragraph('3. A Colheita (Harvest):', style='List Number')
    document.add_paragraph('No final do ciclo, cria-se o registo de Colheita. Este é o momento em que a produção se torna "Inventário", registando quantidade e qualidade (calibre, brix).', style='List Continue')

    document.add_heading('B. O Mercado e Logística', level=2)
    document.add_paragraph('1. Ordens de Mercado (MarketplaceOrder):', style='List Number')
    document.add_paragraph('O Produtor pode criar uma oferta de venda (SELL) ligada a uma Colheita. O Retalhista pode criar um pedido de compra (BUY).', style='List Continue')
    
    document.add_paragraph('2. Transação:', style='List Number')
    document.add_paragraph('Quando um negócio é fechado, a ordem passa a estado APPROVED.', style='List Continue')

    document.add_paragraph('3. Transporte:', style='List Number')
    document.add_paragraph('A ordem contém campos para seguir a logística. Sensores IoT no camião podem enviar dados que ficam guardados no campo transport_sensor_data.', style='List Continue')

    document.add_heading('C. O Papel da Blockchain', level=2)
    document.add_paragraph('A Blockchain não substitui a base de dados principal, funciona como um "Cartório Digital".')
    document.add_paragraph('Sempre que ocorre um evento crítico (ex: Criação de Colheita, Saída para Transporte), o sistema cria um BlockchainBlock.', style='List Bullet')
    document.add_paragraph('O que é guardado: Um "snapshot" dos dados naquele momento (JSON) + uma assinatura criptográfica (hash).', style='List Bullet')
    document.add_paragraph('Para que serve: Garante ao consumidor final que os dados mostrados no QR Code são verdadeiros e não foram alterados posteriormente.', style='List Bullet')

    # Section 4
    document.add_heading('4. Resumo para Apresentação', level=1)
    document.add_paragraph('Estrutura Relacional: Garante a integridade dos dados operacionais.', style='List Bullet')
    document.add_paragraph('Modularidade: Os eventos são modulares, permitindo adicionar novos tipos de sensores.', style='List Bullet')
    document.add_paragraph('Camada de Confiança: A tabela BlockchainBlock corre em paralelo para "selar" a história do produto.', style='List Bullet')

    document.save('Retail_Eureka_DB_Schema_Final.docx')
    print("Document created successfully.")

if __name__ == "__main__":
    create_document()
